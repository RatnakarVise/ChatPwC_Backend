from pathlib import Path
from docx import Document


def create_ts_docx(ts_text: str, output_path: Path) -> None:
    """
    Very simple DOCX generator:
    - Splits by lines
    - Lines starting with '## ' become headings
    - Others become paragraphs
    """
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()

    doc.add_heading("Technical Specification", level=1)

    for line in ts_text.splitlines():
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.strip():
            doc.add_paragraph(line)

    doc.save(str(output_path))
